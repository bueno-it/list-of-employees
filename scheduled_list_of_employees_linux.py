"""
scheduled_list_of_employees_linux.py
Linux version of the Windows scheduled_list_of_employees.py.
Scheduled via cron job — runs after each shift starts.

Setup (first time):
    python3 scheduled_list_of_employees_linux.py --setup

Manual run:
    python3 scheduled_list_of_employees_linux.py

View log:
    tail -f ~/scheduled_list_of_employees.log

Cron examples (edit with: crontab -e):
    0 7    * * * python3 /home/masterlink/Downloads/scheduled_list_of_employees_linux.py
    30 15  * * * python3 /home/masterlink/Downloads/scheduled_list_of_employees_linux.py
    30 23  * * * python3 /home/masterlink/Downloads/scheduled_list_of_employees_linux.py
"""

import os
import sys
import glob
import shutil
import subprocess
import logging
import argparse
from datetime import datetime
from zoneinfo import ZoneInfo

# ──────────────────────────────────────────────────────────────────────────────
# ⚙️  SETTINGS — fill in before running
# ──────────────────────────────────────────────────────────────────────────────
EMAIL    = os.environ.get("EMAIL", "")    # ← full access account email
PASSWORD = os.environ.get("PASSWORD", "")    # ← full access account password

# Path to OneDrive folder via symlink in Documents
EXCEL_PATH = os.path.join(
    os.path.expanduser("~"),
    "Documents",
    "ListOfEmployeesFilesL",
    "list_of_employees.docx"
)

# ensure EXCEL dir exists without failing on symlinks
def ensure_excel_dir():
    parent = os.path.dirname(EXCEL_PATH)
    if os.path.islink(parent) or os.path.isdir(parent):
        return True
    try:
        os.makedirs(parent)
        return True
    except Exception as e:
        log(f"ERROR creating dir: {e}")
        return False

SCRIPT_PATH = os.path.abspath(__file__)

# Cron schedules — 1 hour after each shift starts
CRON_ENTRIES = [
    "0 7 * * *",    # Shift 1 — 07:00
    "30 15 * * *",  # Shift 2 — 15:30
    "30 23 * * *",  # Shift 3 — 23:30
]
# ──────────────────────────────────────────────────────────────────────────────

LOG_PATH = os.path.join(os.path.expanduser("~"), "scheduled_list_of_employees.log")

logging.basicConfig(
    filename=LOG_PATH,
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%d/%m/%Y %H:%M:%S"
)

def log(msg):
    logging.info(msg)
    print(msg)

# ──────────────────────────────────────────────────────────────────────────────
# PATHS
# ──────────────────────────────────────────────────────────────────────────────
BROWSERS_PATH   = os.path.join(os.path.expanduser("~"), ".holiday_balance_browsers")
SETUP_DONE_FLAG = os.path.join(os.path.expanduser("~"), ".holiday_balance_setup_done")


def find_chromium_exe():
    """Find chromium — checks snap, system and playwright paths."""
    # check playwright browsers path first
    for pattern in [
        os.path.join(BROWSERS_PATH, "**", "chrome-headless-shell"),
        os.path.join(BROWSERS_PATH, "**", "chrome"),
        os.path.join(BROWSERS_PATH, "**", "chromium"),
    ]:
        matches = [m for m in glob.glob(pattern, recursive=True) if os.path.isfile(m)]
        if matches:
            return matches[0]

    # fallback to system/snap chromium
    for path in ["/snap/bin/chromium", "/usr/bin/chromium", "/usr/bin/chromium-browser"]:
        if os.path.isfile(path):
            return path

    for name in ["chromium", "chromium-browser", "google-chrome"]:
        found = shutil.which(name)
        if found:
            return found

    return None


# ──────────────────────────────────────────────────────────────────────────────
# SETUP
# ──────────────────────────────────────────────────────────────────────────────
def run_setup() -> bool:
    python = sys.executable
    log(f"Python: {python}")
    os.makedirs(BROWSERS_PATH, exist_ok=True)
    env = os.environ.copy()
    env["PLAYWRIGHT_BROWSERS_PATH"] = BROWSERS_PATH

    # system deps already installed via Dockerfile
    log("System dependencies pre-installed in Docker image.")

    for pkg in ["playwright", "openpyxl", "python-docx"]:
        log(f"Installing {pkg}...")
        r = subprocess.run(
            [python, "-m", "pip", "install", pkg, "--quiet", "--break-system-packages"],
            capture_output=True, text=True, env=env,
            encoding="utf-8", errors="replace"
        )
        if r.returncode != 0:
            log(f"ERROR installing {pkg}: {(r.stderr or r.stdout).strip()}")
            return False
        log(f"{pkg} installed!")

    # use system chromium if available
    chromium = find_chromium_exe()
    if chromium:
        log(f"Using system Chromium: {chromium}")
    else:
        log("Downloading Chromium via Playwright...")
        playwright_cli = shutil.which("playwright")
        cmd = [playwright_cli, "install", "chromium"] if playwright_cli else \
              [python, "-m", "playwright", "install", "chromium"]
        r = subprocess.run(cmd, capture_output=True, text=True, env=env,
                           encoding="utf-8", errors="replace")
        if r.returncode != 0:
            log(f"ERROR: {(r.stderr or r.stdout).strip()}")
            return False
        log("Chromium installed!")

    with open(SETUP_DONE_FLAG, "w") as f:
        f.write("done")

    log("Setup complete!")
    return True


def setup_needed():
    return not os.path.exists(SETUP_DONE_FLAG) or not find_chromium_exe()


# ──────────────────────────────────────────────────────────────────────────────
# CRON
# ──────────────────────────────────────────────────────────────────────────────
def setup_cron():
    python = sys.executable
    result = subprocess.run(["crontab", "-l"], capture_output=True, text=True)
    existing = result.stdout if result.returncode == 0 else ""

    # remove old entries for this script
    lines = [l for l in existing.splitlines() if SCRIPT_PATH not in l and l.strip()]

    # add new entries for each shift
    for schedule in CRON_ENTRIES:
        entry = f"{schedule} {python} {SCRIPT_PATH} >> {LOG_PATH} 2>&1"
        lines.append(entry)
        log(f"Cron registered: {entry}")

    proc = subprocess.run(["crontab", "-"], input="\n".join(lines) + "\n",
                          text=True, capture_output=True)
    if proc.returncode == 0:
        log("All shift cron jobs registered!")
        log("Shift 1: 07:00 daily")
        log("Shift 2: 15:30 daily")
        log("Shift 3: 23:30 daily")
    else:
        log(f"ERROR setting up cron: {proc.stderr}")


def remove_cron():
    result = subprocess.run(["crontab", "-l"], capture_output=True, text=True)
    if result.returncode != 0:
        print("No crontab found.")
        return
    lines = [l for l in result.stdout.splitlines() if SCRIPT_PATH not in l and l.strip()]
    subprocess.run(["crontab", "-"], input="\n".join(lines) + "\n", text=True)
    log("Cron jobs removed.")


# ──────────────────────────────────────────────────────────────────────────────
# MAIN LOGIC — identical to Windows version, only launch args differ
# ──────────────────────────────────────────────────────────────────────────────
def run() -> bool:
    log("=" * 50)
    log("Starting List of Employees sync...")
    log("=" * 50)

    os.makedirs(BROWSERS_PATH, exist_ok=True)
    os.environ["PLAYWRIGHT_BROWSERS_PATH"] = BROWSERS_PATH

    if setup_needed():
        log("Dependencies missing — running setup...")
        if not run_setup():
            log("Setup failed. Aborting.")
            return False

    chromium_exe = find_chromium_exe()
    if not chromium_exe:
        log("ERROR: Chromium not found after setup.")
        return False

    log(f"Chromium: {chromium_exe}")

    from playwright.sync_api import sync_playwright
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    today       = datetime.now(ZoneInfo("Europe/Dublin")).strftime("%d/%m/%Y")
    last_update = datetime.now(ZoneInfo("Europe/Dublin")).strftime("%d/%m/%Y %H:%M")
    log(f"Fetching roll call for {today}...")

    import tempfile, shutil as _shutil
    tmp_profile = tempfile.mkdtemp(prefix="pw_profile_")

    try:
        with sync_playwright() as pl:
            context = pl.chromium.launch_persistent_context(
                user_data_dir=tmp_profile,
                headless=True,
                executable_path=chromium_exe,
                ignore_https_errors=True,
                # Linux-specific args
                args=[
                    "--no-sandbox",
                    "--disable-setuid-sandbox",
                    "--disable-dev-shm-usage",
                    "--disable-gpu",
                ]
            )
            page = context.new_page()

            # login — identical to Windows
            page.goto("https://www.mytimepoint.ie")
            page.fill("#UserName", EMAIL)
            page.fill("#Password", PASSWORD)
            try:
                page.click("button[type='submit'], input[type='submit']")
            except Exception:
                page.press("#Password", "Enter")

            page.wait_for_load_state("networkidle")
            log("OK: Logged in successfully")

            # navigate to establish full session — identical to Windows
            page.goto("https://www.mytimepoint.ie/Home/Index?siteLogin=True")
            page.wait_for_load_state("networkidle")

            # API call — identical to Windows
            response = page.evaluate(f"""
            async () => {{
                const res = await fetch(
                    'https://www.mytimepoint.ie/api/RollCall/GetCurrentRollCallsForSite' +
                    '?siteId=1677&selectedDate={today}&groupName=ClockedDepartment&departmentIds=&costCentreIds=',
                    {{
                        method: 'GET',
                        headers: {{
                            'accept': 'application/json, text/javascript, */*; q=0.01',
                            'accept-language': 'en-US,en;q=0.5',
                            'x-requested-with': 'XMLHttpRequest',
                            'sec-fetch-dest': 'empty',
                            'sec-fetch-mode': 'cors',
                            'sec-fetch-site': 'same-origin'
                        }},
                        credentials: 'include'
                    }}
                );
                if (!res.ok) return {{ errorr: true, status: res.status }};
                return await res.json();
            }}
            """)

            if isinstance(response, dict) and response.get("errorr"):
                log(f"API ERROR: {response.get('status')}")
                context.close()
                return False

            # process data — identical to Windows
            all_employees = []
            if isinstance(response, list):
                for group in response:
                    for emp in group.get("RollCalls", []):
                        if emp.get("IsGroupRow"):
                            continue
                        all_employees.append({
                            "StaffNumber":    emp.get("StaffNumber", ""),
                            "Name":           emp.get("Name", ""),
                            "DepartmentName": emp.get("DepartmentName", ""),
                            "Status":         emp.get("Status", ""),
                        })

            # filter only In and On-Break — identical to Windows
            filtered = [e for e in all_employees if e["Status"] in ("In", "On-Break")]
            log(f"PROCESSING {len(all_employees)} total — {len(filtered)} In/On-Break employees...")

            # build Word document grouped by department
            from docx import Document as DocxDocument
            from docx.shared import Pt, Inches, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from collections import defaultdict
            import os as _os

            ensure_excel_dir()
            doc = DocxDocument()

            # page margins
            section = doc.sections[0]
            section.top_margin    = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin   = Cm(2.0)
            section.right_margin  = Cm(2.0)

            # logo
            script_dir = _os.path.dirname(_os.path.abspath(__file__))
            logo_path  = _os.path.join(script_dir, "masterlink_logo.png")
            if _os.path.exists(logo_path):
                logo_p = doc.add_paragraph()
                logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_logo = logo_p.add_run()
                run_logo.add_picture(logo_path, width=Inches(3.0))
                doc.add_paragraph()

            # title
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run("List of Employees")
            run.bold = True
            run.font.size = Pt(16)

            # date
            date_p = doc.add_paragraph()
            date_p.add_run(f"Date: {today}").bold = True
            doc.add_paragraph()

            # group by department
            dept_groups = defaultdict(list)
            for emp in filtered:
                dept_groups[emp["DepartmentName"]].append(emp)

            # totals
            total_in    = sum(1 for e in filtered if e["Status"] == "In")
            total_break = sum(1 for e in filtered if e["Status"] == "On-Break")

            # table — 4 columns
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr = table.rows[0].cells
            for i, h in enumerate(["Staff Number", "Name", "Department", "Status"]):
                p = hdr[i].paragraphs[0]
                run_h = p.add_run(h)
                run_h.bold = True
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "2E75B6")
                shd.set(qn("w:val"), "clear")
                hdr[i]._tc.get_or_add_tcPr().append(shd)
                run_h.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            widths = [Inches(1.0), Inches(2.5), Inches(2.0), Inches(1.0)]
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    cell.width = widths[i]

            for dept_name, emps in dept_groups.items():
                short_name = dept_name.split("-", 1)[-1].strip() if "-" in dept_name else dept_name

                # department group header row
                dept_row = table.add_row()
                dept_cell = dept_row.cells[0]
                dept_cell.merge(dept_row.cells[3])
                dept_cell.text = f"{short_name}  ({len(emps)})"
                dept_cell.paragraphs[0].runs[0].bold = True
                shd2 = OxmlElement("w:shd")
                shd2.set(qn("w:fill"), "D9D9D9")
                shd2.set(qn("w:val"), "clear")
                dept_cell._tc.get_or_add_tcPr().append(shd2)

                for emp in emps:
                    row_cells = table.add_row().cells
                    for ci, val in enumerate([str(emp["StaffNumber"]), emp["Name"], short_name]):
                        run_c = row_cells[ci].paragraphs[0].add_run(val)
                        run_c.font.size = Pt(8)
                    status_text = "In" if emp["Status"] == "In" else "On-Break"
                    status_run = row_cells[3].paragraphs[0].add_run(status_text)
                    status_run.bold = True
                    status_run.font.size = Pt(8)

            # totals row
            total_row = table.add_row()
            total_cell = total_row.cells[0]
            total_cell.merge(total_row.cells[3])
            p_tot = total_cell.paragraphs[0]
            p_tot.add_run(f"Total: {len(filtered)}   ").bold = True
            p_tot.add_run(f"In: {total_in}   ").bold = True
            p_tot.add_run(f"On-Break: {total_break}").bold = True

            doc.add_paragraph()
            gen = doc.add_paragraph(f"Generated on {today}")
            gen.runs[0].italic = True

            doc.save(EXCEL_PATH)
            log("OK: Document updated successfully!")
            log(f"Path: {EXCEL_PATH}")
            upload_to_onedrive(EXCEL_PATH)
            context.close()
            return True

    except Exception as e:
        log(f"ERROR: {e}")
        return False
    finally:
        _shutil.rmtree(tmp_profile, ignore_errors=True)




def upload_to_onedrive(file_path: str) -> bool:
    """Upload file to OneDrive via Microsoft Graph API."""
    import urllib.request
    import urllib.parse
    import json

    client_id     = os.environ.get("GRAPH_CLIENT_ID", "")
    tenant_id     = os.environ.get("GRAPH_TENANT_ID", "")
    client_secret = os.environ.get("GRAPH_CLIENT_SECRET", "")
    user          = os.environ.get("ONEDRIVE_USER", "")

    if not all([client_id, tenant_id, client_secret, user]):
        log("ERROR: Missing Graph API credentials in environment.")
        return False

    # Get access token
    token_url  = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
    token_data = urllib.parse.urlencode({
        "grant_type":    "client_credentials",
        "client_id":     client_id,
        "client_secret": client_secret,
        "scope":         "https://graph.microsoft.com/.default",
    }).encode()

    req = urllib.request.Request(token_url, data=token_data, method="POST")
    with urllib.request.urlopen(req) as resp:
        token = json.loads(resp.read())["access_token"]

    # Upload file
    filename     = os.path.basename(file_path)
    folder       = "ListOfEmployeesFilesL"
    upload_url   = f"https://graph.microsoft.com/v1.0/users/{user}/drive/root:/{folder}/{filename}:/content"

    with open(file_path, "rb") as f:
        file_data = f.read()

    req = urllib.request.Request(
        upload_url,
        data=file_data,
        method="PUT",
        headers={
            "Authorization": f"Bearer {token}",
            "Content-Type":  "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
    )
    with urllib.request.urlopen(req) as resp:
        result = json.loads(resp.read())
        log(f"OK: Uploaded to OneDrive — {result.get('name')} ({result.get('size')} bytes)")
    return True

# ──────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="List of Employees Sync — Linux")
    parser.add_argument("--setup",       action="store_true", help="Install dependencies and setup cron")
    parser.add_argument("--setup-cron",  action="store_true", help="Setup cron jobs only")
    parser.add_argument("--remove-cron", action="store_true", help="Remove cron jobs")
    args = parser.parse_args()

    if args.setup:
        run_setup()
        setup_cron()
    elif args.setup_cron:
        setup_cron()
    elif args.remove_cron:
        remove_cron()
    else:
        success = run()
        log("Finished — success." if success else "Finished — with errors.")
        sys.exit(0 if success else 1)